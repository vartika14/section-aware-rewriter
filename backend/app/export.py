"""Turn a list of sections back into a real .docx file — the reverse of
parsing.py. Each section becomes a "Heading 1" paragraph plus body text,
except the untitled opening text, which gets no heading at all.
"""

from io import BytesIO

from docx import Document

from .parsing import PREAMBLE_HEADING, Section


def build_docx(sections: list[Section]) -> bytes:
    document = Document()

    for section in sections:
        # The opening text (before any real heading) gets no heading style,
        # so when we read it back in, it's correctly recognized as the
        # opening text again — not as its own numbered section.
        if section.heading != PREAMBLE_HEADING:
            document.add_paragraph(section.heading, style="Heading 1")
        for paragraph in section.text.split("\n\n"):
            document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
