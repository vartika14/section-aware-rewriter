"""Turn a .docx into the sections the rest of the app addresses by id.

.docx only, deliberately. Word carries real heading styles, so section
boundaries are read rather than guessed. PDF would mean inferring structure
from font sizes and positions — a far larger problem for no extra credit.

When a document has no heading styles at all we still return something usable,
but we say so, so the UI can tell the user the split is a guess.
"""

from io import BytesIO

from docx import Document
from pydantic import BaseModel

PREAMBLE_HEADING = "(untitled opening)"


class UnparseableDocument(ValueError):
    """The upload was not a readable .docx, or contained no text."""


class Section(BaseModel):
    id: str
    heading: str
    text: str


class ParsedDocument(BaseModel):
    sections: list[Section]
    headings_detected: bool


def parse_docx(data: bytes) -> ParsedDocument:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:  # noqa: BLE001 — boundary: the bytes are untrusted
        raise UnparseableDocument("That file could not be read as a .docx.") from exc

    paragraphs = [(p.style.name, p.text.strip()) for p in document.paragraphs]

    if not any(text for _, text in paragraphs):
        raise UnparseableDocument("That document contains no text.")

    headings_detected = any(
        style.startswith("Heading") for style, text in paragraphs if text
    )
    pairs = (
        _split_on_headings(paragraphs)
        if headings_detected
        else _split_on_blank_lines(paragraphs)
    )

    sections: list[Section] = []
    n = 1
    for heading, text in pairs:
        if heading == PREAMBLE_HEADING:
            # A fixed id, outside the numbering sequence — so it never shifts
            # what a real section's number means. Only ever produced by the
            # heading-styled path; the blank-line fallback has no preamble
            # concept, since each block's own first line is its heading.
            sections.append(Section(id="preamble", heading=heading, text=text))
        else:
            sections.append(Section(id=f"s{n}", heading=heading, text=text))
            n += 1

    return ParsedDocument(sections=sections, headings_detected=headings_detected)


def _split_on_headings(paragraphs: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Start a new section at every heading-styled paragraph."""
    out: list[tuple[str, str]] = []
    heading: str | None = None
    body: list[str] = []

    for style, text in paragraphs:
        if style.startswith("Heading") and text:
            if heading is not None or body:
                out.append((heading or PREAMBLE_HEADING, "\n\n".join(body)))
            heading, body = text, []
        elif text:
            body.append(text)

    if heading is not None or body:
        out.append((heading or PREAMBLE_HEADING, "\n\n".join(body)))
    return out


def _split_on_blank_lines(paragraphs: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Fallback: treat blank-line-separated blocks as sections, and the first
    line of each block as its heading."""
    blocks: list[list[str]] = []
    current: list[str] = []

    for _, text in paragraphs:
        if text:
            current.append(text)
        elif current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)

    return [(block[0], "\n\n".join(block[1:])) for block in blocks]
