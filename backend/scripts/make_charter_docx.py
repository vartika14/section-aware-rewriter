"""Generate the third sample document — a project charter. A rollout date in
one section, a scope boundary in another: neither is money, neither is a
policy threshold, proving the design doesn't lean on either shape.

  * §2 scopes the pilot to a single warehouse.
  * §3 commits to a go-live date that assumes that scope.
  * §4 names the single stakeholder group being trained, tied to the pilot's
    single-warehouse scope.
"""

from pathlib import Path

from docx import Document

OUTPUT = Path(__file__).resolve().parent.parent / "sample" / "data-platform-charter.docx"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1. Background",
        ["This charter authorises a pilot of the new inventory tracking system."],
    ),
    (
        "2. Scope",
        [
            "The pilot covers the Rotterdam warehouse only. Other sites are "
            "explicitly out of scope for this phase.",
        ],
    ),
    (
        "3. Timeline",
        [
            "Go-live for the Rotterdam warehouse is targeted for the first week "
            "of November, assuming the scope in section 2 holds.",
        ],
    ),
    (
        "4. Training",
        [
            "The Rotterdam warehouse operations team will receive training in "
            "the two weeks before go-live. No other site's staff are scheduled.",
        ],
    ),
]


def main() -> None:
    document = Document()
    document.add_paragraph("Charter: Inventory Tracking Pilot")
    for heading, paragraphs in SECTIONS:
        document.add_paragraph(heading, style="Heading 1")
        for text in paragraphs:
            document.add_paragraph(text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
