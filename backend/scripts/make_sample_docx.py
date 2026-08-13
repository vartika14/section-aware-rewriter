"""Generate the sample proposal used for development and the demo.

Run:  python -m scripts.make_sample_docx      (from backend/)

This document is invented — no client material. It is built so the conflicts are
real rather than decorative. The load-bearing dependencies are:

  * §4 prices a fixed fee "for the scope set out in section 2" — it cites §2 by
    reference, so any change to §2 changes what the fee covers.
  * §4 caps stakeholder interviews at twelve — a hard number another section can
    contradict, and one the document can resolve on its own.
  * §4 invoices in three instalments, coupled to §3's three phases.
  * §1 promises delivery "within the quarter".
  * §5 excludes implementation work, a line §2 only gestures at.

Five sections, six dependencies. Editing any one paragraph can break another.
"""

from pathlib import Path

from docx import Document

OUTPUT = Path(__file__).resolve().parent.parent / "sample" / "meridian-proposal.docx"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1. Executive Summary",
        [
            "Meridian Retail BV has asked us to review its customer data landscape "
            "and recommend a path to consolidation.",
            "We will assess the current state, identify the principal risks, and "
            "deliver a recommendation the leadership team can act on within the quarter.",
        ],
    ),
    (
        "2. Scope of Work",
        [
            "Our work will cover Meridian's customer-facing data systems. We will "
            "review current data flows, interview key stakeholders, and produce "
            "findings and recommendations.",
            "The engagement is advisory; implementation is out of scope.",
        ],
    ),
    (
        "3. Approach and Timeline",
        [
            "The engagement runs over eight weeks in three phases: Discovery "
            "(weeks 1-3), Analysis (weeks 4-6), and Recommendation (weeks 7-8).",
            "A single steering checkpoint is held at the close of each phase.",
        ],
    ),
    (
        "4. Fees and Payment",
        [
            "A fixed fee of EUR 48,000 covers the engagement in full, invoiced in "
            "three instalments at the close of each phase.",
            "The fee assumes the scope set out in section 2 and no more than twelve "
            "stakeholder interviews.",
        ],
    ),
    (
        "5. Assumptions and Exclusions",
        [
            "We assume Meridian provides access to systems documentation within one "
            "week of kickoff.",
            "Excluded: data migration, tooling procurement, and any hands-on "
            "implementation work.",
        ],
    ),
]


def main() -> None:
    document = Document()
    document.add_paragraph(
        "Proposal: Customer Data Platform Review — Meridian Retail BV"
    )

    for heading, paragraphs in SECTIONS:
        document.add_paragraph(heading, style="Heading 1")
        for text in paragraphs:
            document.add_paragraph(text)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
