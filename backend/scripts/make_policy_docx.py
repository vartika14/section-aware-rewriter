"""Generate the second sample document — an internal policy, deliberately with
no money vocabulary, so the interrupt policy is proven on a domain the old
regex-based commitment check would have missed.

  * §2 defines "regular remote work" as up to three days per week.
  * §3 requires manager approval for anything beyond that definition.
  * §4 sets a documentation deadline of five business days after any change.

Editing §2's definition changes what §3's approval rule is measured against.
"""

from pathlib import Path

from docx import Document

OUTPUT = Path(__file__).resolve().parent.parent / "sample" / "remote-work-policy.docx"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1. Purpose",
        ["This policy sets out how remote work is arranged and approved."],
    ),
    (
        "2. Definitions",
        [
            "Regular remote work means working from a non-office location up to "
            "three days per week, on a recurring basis.",
        ],
    ),
    (
        "3. Approval",
        [
            "Regular remote work, as defined in section 2, requires no separate "
            "approval beyond the employee's manager confirming the arrangement in "
            "writing. Anything beyond that definition requires HR sign-off.",
        ],
    ),
    (
        "4. Recordkeeping",
        [
            "Any change to a remote work arrangement must be documented within "
            "five business days of the change taking effect.",
        ],
    ),
]


def main() -> None:
    document = Document()
    document.add_paragraph("Remote Work Policy — Internal")
    for heading, paragraphs in SECTIONS:
        document.add_paragraph(heading, style="Heading 1")
        for text in paragraphs:
            document.add_paragraph(text)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
